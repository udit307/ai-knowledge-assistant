from docx import Document
from openpyxl import Workbook
import os


EXPORT_FOLDER = "exports"

os.makedirs(EXPORT_FOLDER, exist_ok=True)


def export_to_word(question, answer, citations):

    doc = Document()

    doc.add_heading("AI Knowledge Assistant Response", level=1)

    doc.add_heading("Question", level=2)
    doc.add_paragraph(question)

    doc.add_heading("Answer", level=2)
    doc.add_paragraph(answer)

    doc.add_heading("Citations", level=2)

    for citation in citations:
        doc.add_paragraph(citation)

    file_path = os.path.join(
        EXPORT_FOLDER,
        "response.docx"
    )

    doc.save(file_path)

    return file_path

def export_to_excel(question, answer, citations):

    wb = Workbook()

    ws = wb.active

    ws.title = "AI Response"

    ws.append(["Question", question])
    ws.append(["Answer", answer])

    ws.append([])

    ws.append(["Citations"])

    for citation in citations:
        ws.append([citation])

    file_path = os.path.join(
        EXPORT_FOLDER,
        "response.xlsx"
    )

    wb.save(file_path)

    return file_path